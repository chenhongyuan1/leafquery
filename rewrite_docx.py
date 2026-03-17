from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc = Document()
    
    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run('LeafQuery 智农病虫害预警系统\n产品说明书  /  Product Manual')
    run_title.bold = True
    run_title.font.size = Pt(18)
    run_title.font.name = 'Microsoft YaHei'

    p_sub = doc.add_paragraph('版本：V1.5    发布日期：2026年3月    面向对象：竞赛评审 / 研发团队')
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.add_run('\n技术栈：Vue 3 + Spring Boot 3 + Python Flask + MyBatis + 和风天气 JWT API')

    doc.add_heading('◆ 一、产品概述与核心价值', level=1)
    doc.add_heading('▸ 1.1 项目背景', level=2)
    doc.add_paragraph('我国是世界上最大的农业生产国，粮食安全事关国家根本。水稻、玉米、小麦三大主粮的病虫害每年造成巨大减产损失。目前农业病虫害的防控存在响应慢、专家资源紧缺等痛点。"精准农业"和"数字农业"是国家农业现代化的战略方向。本项目 LeafQuery（"叶问"）通过将 AI 图像识别、气象大数据、规则引擎与知识库深度融合，构建了一套面向大众农户的轻量化、移动端优先的智慧农业辅助决策系统。')
    
    doc.add_heading('▸ 1.2 产品定位', level=2)
    doc.add_paragraph('LeafQuery 是一款主打"病虫害识别 + 气象关联趋势预警"双核功能的农业 AI 移动与多端应用。目标是让缺乏植保背景的普通农民，在 30 秒内完成"拍照识图 → 查趋势 → 获建议"的闭环。\n目标用户群：华北、东北、长江中下游、华南、西南、西北六大农业生态区的水粮作物种植户及农技推广员。')
    
    doc.add_heading('▸ 1.3 六大农业生态区自动覆盖', level=2)
    doc.add_paragraph('系统通过调用和风天气城市搜索 API（GeoAPI），通过用户模糊搜索即可自动将城市精确定位，并归类至六大农业生态区，使历史气候模型与实际种植环境精确匹配。')

    doc.add_heading('◆ 二、系统技术架构与全栈技术栈', level=1)
    doc.add_heading('▸ 2.1 总体多层架构', level=2)
    doc.add_paragraph('展现层：基于 Vue 3 + Pinia + Vue Router 构建移动端/PC端 SPA 应用，采用 Vite 极速构建，结合 TailwindCSS 原子化样式与 Web Audio/Canvas 渲染高级物理特效。')
    doc.add_paragraph('业务逻辑层：基于 Spring Boot 3 (Java 17)，提供强健的 RESTful API，涵盖农场、预测、社区交互等核心功能引擎。')
    doc.add_paragraph('数据与 AI 层：Python Flask 微服务负责深度学习 CV 模型图像病理分析；底层数据持久化采用 MySQL 8 与 MyBatis 框架；气象业务高度解耦集成和风天气云端 API。')
    
    doc.add_heading('▸ 2.2 核心架构与工程亮点', level=2)
    doc.add_paragraph('• 和风天气 JWT 高级强鉴权：采用 Ed25519 (EdDSA) 算法构建动态令牌，私钥硬隔离，前端零接触。\n• 响应式状态流与持久化策略：整合 Pinia Store + LocalStorage 提供平滑离线/断网容错体验；长文本 Base64 及病害结果快照全部在云端数据库长期存储，前端状态心跳级同步。\n• 优雅容错降级：首创针对失效图片链接的前端兜底策略，历史遗留脏图片链接自动静默回退并渲染为绿色有机 Emoji (🌿/🦠/🐛)，护航全局 UI 一致性。')

    doc.add_page_break()

    doc.add_heading('◆ 三、移动端核心页面逻辑与主要功能', level=1)
    doc.add_paragraph('考虑到终端农户在田间地头的实际使用场景，系统遵循 Mobile-First 设计理念。系统模块以四个核心页面贯穿农作全流程：')

    doc.add_heading('▸ 3.1 识别主页面（Identification.vue）', level=2)
    doc.add_paragraph('作为 LeafQuery 的主入口与首屏，承载直接且高频的诊断诉求，并大幅运用拟真沉浸式视控手段。')
    doc.add_paragraph('• AI 极速病害判别流：用户摄入或上传病叶图片后流转至 Python 微服务推理，极秒内获取高置信度推断与初步防护建议清单。')
    doc.add_paragraph('• 智能沉浸式语音录入 (Web Audio/Canvas API)：将传统呆板的录音按钮置换为人声驱动的“Organic AI Orb”全屏交互球。结合分贝与能量监测算法控制多阶贝塞尔曲线及渐变胶片光晕，呈现随真实语音共振流变的高级视觉特效；同时底座集成超大容错率的防误触唤醒，适配多省市方言录音下钻推理引擎。')
    doc.add_paragraph('• Toast 弱侵入感知：全面废弃粗暴阻断的主流 Alert 交互弹窗，所有权限受限或解析挂起的网络抖动，一律采用色彩梯度级 (Info/Warning/Error) 柔性 Toast 滑块处理反馈，保持核心工作流的顺享性。')

    doc.add_heading('▸ 3.2 发现页面（Discovery.vue）', level=2)
    doc.add_paragraph('定位为全链路双向联通的农业内容矩阵。本次迭代抛却所有 Mock 流，全面下沉由 MyBatis 及 Spring Boot 云端直供数据，打造动态三农生态中心。')
    doc.add_paragraph('• 多维知识图鉴库（Knowledge）：分类归档病毒、病菌、生理病变等作物病理特征卡片。配有极深的高反差渐变模糊 Modal 浮层提供深层次的用药方解说。')
    doc.add_paragraph('• 农技资讯与前沿热推（News）：自适配布局承载涉农热点政令及用药专刊，提供标签分级及文章点击热度统筹流。')
    doc.add_paragraph('• 专家及同好问答圈（Q&A）：运用 JSON Array 数组巧妙封装极密的多宫格问诊贴文存储，避除重库关联阵列；贴内集成了点星响应与特别区分的绿色专家布道区盖楼功能。')
    doc.add_paragraph('• 全局收藏挂载体系：采用统一化的视觉金星号标的，融合 Pinia 全局跨组件传信，将新闻精选及图鉴珍藏秒切归入数字个人的私库中。')

    doc.add_heading('▸ 3.3 预测页面（Prediction.vue）', level=2)
    doc.add_paragraph('实现“防范于未然”的关键阵地。将冷硬的天文数据转化为农人可实操的行动派盘面。')
    doc.add_paragraph('• 五因子生化加权引擎：串联和风核心气象预测数据流。凭借提取的未来七天湿度阵雨累积及长线均温，喂入引擎获取出百分比制的 先验风险阈分 (Prior Risk Score)。')
    doc.add_paragraph('• 物候期动态拟合：容纳多作物苗期、抽穗到成熟期的环境不同脆弱级别，实现针对目标作伴特定病害（例如稻瘟 / 锈病）的客制化报警参数权重配置。')
    doc.add_paragraph('• 自然语义仪表盘：以红橙黄渐进色卡传达等级危机指示；特制“主驱动因子 (topDrivers)”解析器，以“过去累计暴雨高危”、“即将抽穗”等浅白的人类自然语言告知风险内因，辅助实地操作。')

    doc.add_heading('▸ 3.4 我的页面（Profile.vue & 农场中心）', level=2)
    doc.add_paragraph('集成全生命周期的系统权限、用户核心资产沉淀及跨界面协同管理中心。')
    doc.add_paragraph('• 农场自定义域（My Farm）：调用 GeoAPI 进行全图高精细模糊省市下探检索，将主粮品种与 6 大环境气候地带做关联落库。定位的精准 LocationID 由前端静默接管赋能给后段气象算法。')
    doc.add_paragraph('• 全端云数据库同步与溯源：首页最近扫描的历史追溯图片乃至在发现页点亮的所有收藏卡片资源在此汇聚，凭借 Vue Router 极致的无刷新跳跃以及全组件双向绑定，提供如原生客户端般丝滑的数据管理总览平台。')

    doc.add_page_break()

    doc.add_heading('◆ 四、PC 端高级视觉特效与沉浸式体验', level=1)
    doc.add_paragraph('鉴于 PC 大屏展现的宽容度与超大空间优势，我们在开发桌面端排版时不拘泥于传统的列表陈列，引入了重度 CSS 3D 渲染与动能转化特性，向具有宏大科技感的农技中枢靠齐。')
    doc.add_paragraph('• 沉浸式 3D 柱状曲面卡片轮播 (Inside Cylinder Carousel)：借助 2000px 超宽幅远视距与 140vw 的柱面曲率调和，我们打造了第一人称视角的 3D 内容视轴。伴随卷轴下潜，大屏矩阵卡片会基于透视角自动淡入、悬停、逼近到包围玩家周边，构图具备深度电影感推镜体验。')
    doc.add_paragraph('• Canvas 粒子互斥与全局网格畸变 (Repel & Squeeze)：于 PC 大屏末端的宣发交互地带，摒除呆滞图案铺设。中心按钮升格为具备原生网格引力波的 "GO" 唤醒盘，任何鼠绘逼近都会激活其背部极其凶猛的 2D 微型物理引擎。')
    doc.add_paragraph('• 赛博非线性曲线张力爆发 (Power 12)：摒弃枯燥扁滑的前段调参，深探十二次幂的高数加速函数。通过冗长蓄力克制与末尾瞬时全屏放大的 1.15 倍等比全局缩放（Global Zoom），网格交叉点在触发的一瞬间被洪流般抛开，从而缔造最极致的视觉张力，赋予全网农业应用史上顶级的交互品质感。')

    doc.add_page_break()

    doc.add_heading('◆ 五、未来规划版图（V2 与 V3 蓝图规划）', level=1)
    doc.add_heading('▸ V2 机器学习核心增强', level=2)
    doc.add_paragraph('• XGBoost / LightGBM 集成推断模型：融合更丰沛的地域和植保病历标签打磨复合训练。提供 SHAP 模型透明支持，在提升精准防范度同时不减弱“主因白话输出”的亲民度。')
    doc.add_paragraph('• 态势可视指挥库生成系统：基于全站六大农业模块归档频次聚合构建的大数据下沉渲染，投射到多级 ECharts 下的广域宏观大盘，直挂全国微观农损走势。')

    doc.add_heading('▸ V3 全模态大世界顾问架构', level=2)
    doc.add_paragraph('• RAG 大模型向量赋能问诊：将深邃的病理手册向量化存储。打通大语言决策流 (Doubao/GPT)，供基层农人无缝提问并直接输出含精确剂量指引与下次随访计划在内的 400 字闭环诊疗方案。')

    doc.add_heading('◆ 六、结语与架构实证摘要', level=1)
    doc.add_paragraph('LeafQuery 自始至终拒绝前端花样翻新的伪装与数据 Mock。这是一套满负荷运转的全模态级工业产出：后端安全挂载 Ed25519 的硬核隔离验签，中继提供高度平顺的 Pinia & LocalStorage 冗余离线心跳架构，前沿交互更是采用纯粹的 Canvas 与 Web Audio 原理突破 Web 界限，全方位为国内千万基干农户编织了一把数字化科技巨伞。')

    doc.save('产品说明书.docx')
    print("Document successfully regenerated and saved.")

if __name__ == '__main__':
    main()
