from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc_path = '产品说明书.docx'

try:
    doc = Document(doc_path)
except Exception as e:
    print(f"Error loading document: {e}")
    exit(1)

doc.add_page_break()

# ---------- 第一部分：语音交互与体验 ----------
title1 = doc.add_heading('◆ 智能语音识别与交互体验升级（新增特性）', level=1)

p_voice = doc.add_paragraph('为了提升农户在田间地头实地操作时的便捷性，LeafQuery 系统在原有的图像识别、文本交互基础上，深度结合了前沿的 Web Audio API 与 HTML5 Canvas 技术，首创了农技问答场景下的沉浸式（Immersive）语音交互体验。本次更新主要包含以下核心功能亮点：')

doc.add_heading('1. Organic AI Orb：仿生流体视觉反馈', level=2)
p1 = doc.add_paragraph('在用户开启语音录入时，传统的生硬按键被替换为全屏级“有机流体球（Organic AI Orb）”。系统在背后利用了基于高分贝检测和均方根（RMS）音频能量分析算法。通过精确复刻 7 个控制点的贝塞尔曲线（Bezier Curve）与正余弦波形叠加，流体球在接收到用户发声时会随音量强度产生真实、有规律的外扩震动与形态变化。')
p1.add_run('\n• 视觉渲染：').bold = True
p1.add_run(' 采用纯原生 Canvas 绘制配合多重径向渐变，形成契合系统植保主题的“绿-青”动态融合表现。加上动态高斯模糊 Glow 辉光和混合叠加的 Film Grain（胶片噪点），提供具有生机感（Organic）的高级视觉质感，且在所有主流浏览器中均保持 60fps 的稳定帧率交互。')

doc.add_heading('2. 无缝融合的长按录音与多方言支持', level=2)
p2 = doc.add_paragraph('• 交互方式：')
p2.runs[0].bold = True
p2.add_run(' 用户只需“长按”屏幕下方的超大号麦克风悬浮按钮，即可唤出沉浸式 AI 倾听覆盖层（Overlay）；“松开”手指即刻完成录音提交，极大地简化了复杂环境中的操作流程，防止用户误触。')
p2.add_run('\n• 语音流处理：').bold = True
p2.add_run(' 考虑到农民群体的使用习惯，系统深度接入了多方言支持功能（含粤语、上海话、闽南语、四川话、陕西话等）。前端直接利用 ScriptProcessor 捕获音频 PCM 流并进行重采样，转码为 16kHz 标准 WAV 格式后，经由云端火山引擎大模型实时解析。')

doc.add_heading('3. 全局 Toast 弱侵入式异常消息提示', level=2)
p3 = doc.add_paragraph('鉴于实际使用场景中可能存在的麦克风权限阻止、弱网状态或语音未识别出有效内容，传统的 Alert 弹窗或错误对话内容极其破坏交互连贯性。本次重构彻底引入了一套顶部滑出的弱侵入式 Toast 通知系统：')
p3.add_run('\n• 智能嗅探：').bold = True
p3.add_run(' 自动拦截后端抛出的无效配置（如 API 密钥挂载失败）、超时、分析异常或“未识别出有效内容”提示，避免无效指令混入问答对话框。')
p3.add_run('\n• 视觉分级：').bold = True
p3.add_run(' 提供 info、warning（琥珀色）、error（红色）多级颜色匹配与自动消失机制，确保异常反馈既被用户感知，又不阻塞核心流程。')

doc.add_heading('4. 历史记录一体化与全平台数据互通', level=2)
p4 = doc.add_paragraph('为了帮助农户回顾溯源此前的病虫害记录，我们：')
p4.add_run('\n• 本地与云端双向同步：').bold = True
p4.add_run(' 最新扫描的病虫害置信度、图片、病情名称及分析会同时存入本地 localStorage 和云端 Database，首页“最近记录”横列可随手翻阅历史卡片。')
p4.add_run('\n• 功能直达互通：').bold = True
p4.add_run(' 首页点击“全部”按钮即可实现 Vue Router 的跨页面平滑跳转，连接至新建的“识别记录”总览页面，方便后续导出甚至咨询专家。')


doc.add_page_break()

# ---------- 第二部分：前端视觉特效 ----------
title2 = doc.add_heading('◆ 前端视觉与物理动效全面升级（PC 端高级特性）', level=1)

p_vis = doc.add_paragraph('在打造沉浸式的现代 Web 体验方面，LeafQuery 针对 PC 平台进行了系统级视觉重构。我们抛弃了传统的平面布局，全面引入了基于原生 CSS 3D 转换与 HTML5 Canvas 的物理级特效，旨在打造极具未来科技感与高级质感的用户连接点。本次主要更新包括以下核心方向：')

doc.add_heading('1. 沉浸式 3D 柱状曲面卡片轮播 (Inside Cylinder Carousel)', level=2)
p5 = doc.add_paragraph('针对农业核心技术与功能的展示模块 "VISION"，团队首创了第一人称视角的 3D 内圆柱面滚轮（Inside Cylinder Track）。')
p5.add_run('\n• 空间透视算法：').bold = True
p5.add_run(' 利用动态计算的 Y 轴旋转角结合 `140vw` 的平缓曲率半径（Radius）以及高达 `2000px` 的视距（Perspective），构建出包围用户的全景式大屏卡片组。')
p5.add_run('\n• 滚动融合驱动：').bold = True
p5.add_run(' 所有卡片的入场、退场、位移及透明度，均深度绑定用户滚动进度（Scroll Progress）。搭配 `+/- 60vh` 的大跨度错落排版（Stagger），实现了"推拉镜头"般的极致纵深沉浸感，有效聚合了海量图文信息。')

doc.add_heading('2. 基于 Canvas 的动力学粒子与网格畸变引擎', level=2)
p6 = doc.add_paragraph('在网页底部的产品下载（CTA）区域，为了强化高频转化场景的交互质感，我们用高度定制的物理变形动效取代了传统的“死白”二维码展示框。')
p6.add_run('\n• 流变式“GO”感应区：').bold = True
p6.add_run(' 下载触点被设计为一个发光的“GO”呼吸圆盘。当用户鼠标悬停时，圆圈将顺滑展开并呈现出其内部暗藏的微观粒子二维码。')
p6.add_run('\n• 引力互斥与空间畸变 (Repel & Squeeze)：').bold = True
p6.add_run(' 底层部署了基于 60fps 重新计算的微型 2D 物理引擎。网格系统的每个相交点都会实时监听“GO”圆盘的状态，并在用户悬停时被强大“排斥力”迅速挤压排开。')

doc.add_heading('3. 极致调教的非线性运动美学', level=2)
p7 = doc.add_paragraph('在这套物理背景网格系统的动效表现上，团队经过反复打磨，引入了极端的加速模拟，彻底消除“线性动画”的干瘪感：')
p7.add_run('\n• 十二次幂 (Power 12) 加速曲线：').bold = True
p7.add_run(' 为了营造出强烈的“张力爆发”感，我们在底层数学运算中应用了高次幂算法。在整个 1000ms 的动画生命周期中，前 80% 的时间被“克制”在近乎静止的蓄力状态，而在结尾几十毫秒内所有元素如火山般瞬间弹开。')
p7.add_run('\n• Global Zoom 全局缩放：').bold = True
p7.add_run(' 网格被挤压推开的同时，底层画布会同步跟进 1.15 倍的等比放大。配合 `cubic-bezier(0.95, 0, 0.05, 1)` 带来的视觉冲击，使得简单的绿色坐标系网格爆发出宏大、富有空间扭曲感的高级科技质感，并完美契合上方 VISION 区的亮绿色彩基调。')

doc.save(doc_path)
print(f"Successfully updated {doc_path} with Voice & Visulizations details.")
